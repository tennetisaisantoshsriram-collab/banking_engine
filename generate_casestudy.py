"""
Generates the IBM-format Case Study DOCX for:
Banking Credit Default Risk & Cross-Sell Engine
Student: T. Sai Santosh Sri Ram
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SHOTS = "C:/Users/sait9/OneDrive/Desktop/banking_engine/docs/screenshots"

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

DARK_BLUE  = (31, 56, 100)
MID_BLUE   = (68, 114, 196)
WHITE      = (255, 255, 255)
LIGHT_FILL = "DCE6F1"
HEAD_FILL  = "1F3864"

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_font(run, size=11, bold=False, color=None, italic=False, name="Calibri"):
    run.font.name   = name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, size=14, color=DARK_BLUE, center=False, bold=True):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return p

def body(text, indent=False, bold=False, italic=False, size=11, center=False):
    p = doc.add_paragraph()
    if indent:  p.paragraph_format.left_indent = Cm(0.5)
    if center:  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, italic=italic)
    return p

def bullet(text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_font(r, size=size)
    return p

def shade_cell(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill)
    tcPr.append(shd)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        shade_cell(cell, HEAD_FILL)
        r = cell.paragraphs[0].runs[0]
        set_font(r, size=10, bold=True, color=WHITE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            r = cell.paragraphs[0].runs[0]
            set_font(r, size=10)
            if ri % 2 == 1:
                shade_cell(cell, LIGHT_FILL)
    return t

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    r = p.add_run(text)
    r.font.name  = "Courier New"
    r.font.size  = Pt(9)
    r.font.color.rgb = RGBColor(20, 20, 20)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F0F0F0")
    pPr.append(shd)
    return p

def add_image(filename, caption=None, width=Inches(5.8)):
    path = f"{SHOTS}/{filename}"
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            c = doc.add_paragraph()
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = c.add_run(caption)
            set_font(r, size=9, italic=True, color=(100, 100, 100))
    else:
        body(f"[Screenshot not found: {filename}]", italic=True)

def sp():
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
heading("IBM CASE STUDY", size=22, color=DARK_BLUE, center=True)
sp()
heading("Banking Credit Default Risk & Cross-Sell Engine", size=16, color=DARK_BLUE, center=True)
sp(); sp()

cover = [
    ("Name",          "T. Sai Santosh Sri Ram"),
    ("Roll No",       "A24126552121"),
    ("Department",    "CSM (AI & ML)"),
    ("College Name",  "Anil Neerukonda Institute of Technology and Sciences"),
    ("Module Name",   "AIML"),
    ("Module Number", "12"),
    ("UG Level",      "UG2"),
]
for label, value in cover:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label:<20}")
    set_font(r1, size=12, bold=True, color=DARK_BLUE)
    r2 = p.add_run(value)
    set_font(r2, size=12)

sp()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("A Case Study and AIML Mini Project Report")
set_font(r, size=13, bold=True, italic=True, color=DARK_BLUE)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════
heading("ABSTRACT", size=14, color=DARK_BLUE)
body(
    "Modern banks and financial institutions are exposed to two fundamental business challenges: "
    "minimizing losses due to loan defaults and maximizing revenue through targeted cross-selling of "
    "financial products. Conventional rule-based risk models fail to capture complex behavioral "
    "patterns embedded in customer financial data, while product recommendations are typically based "
    "on generic segmentation rather than individual transaction history."
)
body(
    "This project proposes a Banking Credit Default Risk and Cross-Sell Engine using Machine Learning. "
    "The system analyzes a dataset of 2,000 banking customers containing demographic, financial, and "
    "product-ownership attributes. A Random Forest Classifier is trained to predict the probability "
    "that a customer will default on their loan. A market-basket-analysis cross-sell engine "
    "simultaneously computes conditional product-affinity scores to recommend the most relevant "
    "financial products to each customer."
)
body(
    "The system further computes SHAP-approximated explainability values, estimates Customer Lifetime "
    "Value (CLV), generates a fraud-anomaly probability score, and includes a macro-economic scenario "
    "simulator. The entire system is deployed as an interactive glassmorphism web dashboard (Aether AI) "
    "built with FastAPI and vanilla JavaScript with secure login and signup functionality."
)
p = doc.add_paragraph()
r = p.add_run("Keywords: ")
set_font(r, size=11, bold=True)
r2 = p.add_run(
    "Credit Default Risk, Cross-Sell Engine, Random Forest, SHAP Explainability, "
    "Customer Lifetime Value, Fraud Detection, FastAPI, Market Basket Analysis, Banking AI, Predictive Analytics."
)
set_font(r2, size=11, italic=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS  (proper 2-column table with page numbers)
# ══════════════════════════════════════════════════════════════════════════════
heading("TABLE OF CONTENTS", size=14, color=DARK_BLUE)
sp()

toc_entries = [
    ("Abstract",                                              "2"),
    ("1.  Introduction",                                      "4"),
    ("2.  Background of the Study",                           "4"),
    ("3.  Problem Statement",                                 "5"),
    ("4.  Objectives",                                        "5"),
    ("5.  Scope of the Project",                              "6"),
    ("6.  Existing System",                                   "6"),
    ("7.  Limitations of Existing System",                    "6"),
    ("8.  Proposed System",                                   "7"),
    ("9.  Advantages of Proposed System",                     "7"),
    ("10. Literature Review",                                  "8"),
    ("11. Case Study",                                         "8"),
    ("    11.1 Business Scenario",                             "8"),
    ("    11.2 Example Customer Analysis",                     "9"),
    ("12. Dataset Description",                                "10"),
    ("13. Feature Description",                                "11"),
    ("    13.1 Engineered Features",                           "11"),
    ("14. Data Preprocessing",                                 "12"),
    ("15. Exploratory Data Analysis",                          "13"),
    ("16. Machine Learning Methodology",                       "13"),
    ("    16.1 Random Forest Classifier",                      "13"),
    ("    16.2 SHAP Explainability Approximation",             "14"),
    ("    16.3 Macro-Economic Scenario Simulation",            "14"),
    ("    16.4 Customer Lifetime Value & Fraud Score",         "15"),
    ("17. System Architecture",                                "15"),
    ("18. Model Training and Evaluation",                      "16"),
    ("19. Risk Classification & Early Warning Mechanism",      "17"),
    ("20. Cross-Sell Recommendation Engine",                   "17"),
    ("21. System Implementation",                              "18"),
    ("22. Web Dashboard — Aether AI",                          "19"),
    ("23. Expected Results",                                   "20"),
    ("24. Advantages",                                         "20"),
    ("25. Limitations",                                        "21"),
    ("26. Future Scope",                                       "21"),
    ("27. Conclusion",                                         "22"),
    ("28. References",                                         "23"),
]

toc_table = doc.add_table(rows=len(toc_entries) + 1, cols=2)
toc_table.style = "Table Grid"
toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
for ci, hdr in enumerate(["Section", "Page"]):
    cell = toc_table.rows[0].cells[ci]
    cell.text = hdr
    shade_cell(cell, HEAD_FILL)
    r = cell.paragraphs[0].runs[0]
    set_font(r, size=11, bold=True, color=WHITE)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Set column widths via first row
from docx.oxml import parse_xml
from docx.oxml.ns import nsmap
tbl = toc_table._tbl
tblGrid = OxmlElement("w:tblGrid")
for w in [9000, 1000]:  # twips
    col = OxmlElement("w:gridCol")
    col.set(qn("w:w"), str(w))
    tblGrid.append(col)
tbl.insert(0, tblGrid)

# Data rows
for ri, (section, page) in enumerate(toc_entries):
    row = toc_table.rows[ri + 1]
    # section name
    cell_sec = row.cells[0]
    cell_sec.text = section
    r = cell_sec.paragraphs[0].runs[0]
    is_sub = section.startswith("    ")
    set_font(r, size=10, bold=(not is_sub), color=DARK_BLUE if not is_sub else (0, 0, 0))
    if ri % 2 == 1:
        shade_cell(cell_sec, LIGHT_FILL)
    # page number
    cell_pg = row.cells[1]
    cell_pg.text = page
    r2 = cell_pg.paragraphs[0].runs[0]
    set_font(r2, size=10, bold=True)
    cell_pg.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if ri % 2 == 1:
        shade_cell(cell_pg, LIGHT_FILL)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
heading("1. INTRODUCTION", size=13, color=DARK_BLUE)
body(
    "The banking and financial services industry handles enormous volumes of lending activity every "
    "year. A loan default occurs when a borrower fails to meet the contractual repayment obligations "
    "of a loan. Defaults result in direct financial losses, increased regulatory capital requirements, "
    "and reduced profitability. Accurately predicting which customers are likely to default before "
    "loan disbursement is one of the most critical applications of machine learning in finance."
)
body(
    "Simultaneously, banks have a strong incentive to cross-sell financial products to existing customers. "
    "A customer who already holds a Checking and Savings account may be a strong candidate for a "
    "Credit Card or Auto Loan. However, recommending the wrong product to the wrong customer wastes "
    "marketing resources and damages customer relationships."
)
body(
    "This project builds a unified Banking Credit Default Risk and Cross-Sell Engine that addresses "
    "both problems simultaneously — presenting results through an interactive Aether AI web dashboard "
    "designed for bank loan officers and relationship managers."
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. BACKGROUND
# ══════════════════════════════════════════════════════════════════════════════
heading("2. BACKGROUND OF THE STUDY", size=13, color=DARK_BLUE)
body(
    "Credit risk assessment has been fundamental to banking since formal lending began. Traditional "
    "FICO scores assign a single numerical value based on payment history and credit utilization. "
    "Machine learning offers the ability to train custom risk models directly on the bank's internal "
    "data, incorporating signals such as income-to-loan ratio and employment stability into a single "
    "predictive model. Ensemble methods such as Random Forest consistently outperform logistic regression "
    "baselines on credit risk datasets."
)
body(
    "Cross-selling in banking has traditionally relied on product bundling and relationship manager "
    "judgment. Data-driven approaches using association rule mining have shown significant improvements "
    "in recommendation relevance and conversion rates."
)
body("Key indicators associated with high default risk include:")
for item in ["Low credit score (below 600)", "High debt-to-income ratio",
             "Short or unstable employment history", "High loan amount relative to income",
             "Low income levels"]:
    bullet(item)

# ══════════════════════════════════════════════════════════════════════════════
# 3. PROBLEM STATEMENT
# ══════════════════════════════════════════════════════════════════════════════
heading("3. PROBLEM STATEMENT", size=13, color=DARK_BLUE)
body(
    "Banks face two interconnected challenges. First, traditional credit risk models rely on static "
    "thresholds and external credit scores that do not leverage the full richness of internal customer "
    "data. A single-number threshold cannot capture the interaction between multiple risk factors — "
    "a customer with a credit score of 600, stable high income, and modest loan amount may be low "
    "risk, while the same score with zero employment and a large loan is very high risk."
)
body(
    "Second, product cross-sell recommendations are often generic, ignoring the customer's existing "
    "product portfolio. A customer who already holds a Mortgage and Investment account may be a poor "
    "target for an Auto Loan but an excellent target for a Credit Card. Neither challenge is currently "
    "addressed with explainability — loan officers need to know why a model flagged a customer as high "
    "risk, not just that it did."
)

# ══════════════════════════════════════════════════════════════════════════════
# 4. OBJECTIVES
# ══════════════════════════════════════════════════════════════════════════════
heading("4. OBJECTIVES", size=13, color=DARK_BLUE)
for i, obj in enumerate([
    "Collect and preprocess a structured banking customer dataset with financial and demographic attributes.",
    "Perform Exploratory Data Analysis to understand default distribution and product ownership.",
    "Train a Random Forest Classifier to predict individual-level loan default probability.",
    "Generate SHAP-approximated feature importance values to explain each prediction.",
    "Compute Customer Lifetime Value (CLV) as a measure of long-term customer worth.",
    "Detect anomalous customer profiles that may indicate fraudulent applications.",
    "Build a market-basket-analysis cross-sell engine to recommend relevant financial products.",
    "Simulate the impact of macro-economic environments on default risk.",
    "Develop an interactive web dashboard for loan officers to analyze customers in real time.",
    "Implement secure user authentication with login and signup functionality.",
    "Deploy the system as a production-ready FastAPI web application.",
], 1):
    body(f"{i}. {obj}")

# ══════════════════════════════════════════════════════════════════════════════
# 5. SCOPE
# ══════════════════════════════════════════════════════════════════════════════
heading("5. SCOPE OF THE PROJECT", size=13, color=DARK_BLUE)
body(
    "The project focuses on supervised binary classification for loan default prediction using a "
    "dataset of 2,000 banking customers. Features include age, income, employment length, credit score, "
    "and loan amount. The cross-sell engine operates over seven product categories: Checking, Savings, "
    "Credit Card, Auto Loan, Mortgage, Personal Loan, and Investment. The project is implemented in "
    "Python using FastAPI and vanilla JavaScript with Chart.js."
)

# ══════════════════════════════════════════════════════════════════════════════
# 6. EXISTING SYSTEM
# ══════════════════════════════════════════════════════════════════════════════
heading("6. EXISTING SYSTEM", size=13, color=DARK_BLUE)
body("In traditional banking, credit risk assessment is performed using:")
for item in [
    "External FICO or CIBIL credit scores from third-party bureaus.",
    "Manual underwriting by loan officers based on income documents.",
    "Simple rule-based systems: IF credit_score < 600 THEN flag as high risk.",
    "Periodic batch processing generating weekly or monthly risk reports.",
    "Generic marketing campaigns for cross-selling based on age or income segment.",
]:
    bullet(item)

# ══════════════════════════════════════════════════════════════════════════════
# 7. LIMITATIONS
# ══════════════════════════════════════════════════════════════════════════════
heading("7. LIMITATIONS OF EXISTING SYSTEM", size=13, color=DARK_BLUE)
for item in [
    "Relies on external scores that may not reflect internal customer behavior.",
    "Rule-based thresholds cannot capture non-linear feature interactions.",
    "Provides no individual-level default probability — only binary classification.",
    "Produces no explanation for why a customer was flagged as risky.",
    "Cross-sell campaigns are generic and not personalized to product portfolios.",
    "Cannot simulate the effect of macro-economic changes on portfolio risk.",
    "No integrated Customer Lifetime Value or fraud detection.",
]:
    bullet(item)

# ══════════════════════════════════════════════════════════════════════════════
# 8. PROPOSED SYSTEM
# ══════════════════════════════════════════════════════════════════════════════
heading("8. PROPOSED SYSTEM", size=13, color=DARK_BLUE)
body("The proposed system pipeline:", bold=True)
for step in [
    "Banking Customer Dataset (2,000 records, 14 features)",
    "↓", "Data Cleaning & Preprocessing", "↓",
    "Feature Engineering (Debt-to-Income Ratio, CLV)", "↓",
    "Train/Test Split (80% / 20%)", "↓",
    "Random Forest Classifier → Default Probability", "↓",
    "SHAP Approximation → Feature Impact Explanation", "↓",
    "Market Basket Analysis → Cross-Sell Affinity Scores", "↓",
    "Fraud Anomaly Detector → Fraud Probability", "↓",
    "Risk Classification: Low / Medium / High", "↓",
    "Macro-Economic Scenario Simulator", "↓",
    "Aether AI Web Dashboard (FastAPI + JavaScript)",
]:
    body(step, indent=True)

# ══════════════════════════════════════════════════════════════════════════════
# 9. ADVANTAGES OF PROPOSED SYSTEM
# ══════════════════════════════════════════════════════════════════════════════
heading("9. ADVANTAGES OF PROPOSED SYSTEM", size=13, color=DARK_BLUE)
for item in [
    "Individual-level default probability score for every customer.",
    "SHAP-approximated explainability — identifies top risk drivers per customer.",
    "Personalized cross-sell recommendations based on actual product ownership.",
    "Customer Lifetime Value estimation for revenue prioritization.",
    "Fraud anomaly score to flag suspicious applications.",
    "Real-time macro-economic scenario simulation (Recession, High Interest, Booming).",
    "Interactive Aether AI glassmorphism dashboard with secure authentication.",
    "Production-ready FastAPI deployment with Render.com support.",
]:
    bullet(item)

# ══════════════════════════════════════════════════════════════════════════════
# 10. LITERATURE REVIEW
# ══════════════════════════════════════════════════════════════════════════════
heading("10. LITERATURE REVIEW", size=13, color=DARK_BLUE)
body(
    "Credit default prediction has been extensively studied using logistic regression, decision trees, "
    "random forests, gradient boosting, and neural networks. Ensemble methods — particularly Random "
    "Forest and XGBoost — consistently achieve the highest AUC-ROC scores on benchmark credit datasets "
    "due to their ability to model non-linear feature interactions and handle class imbalance."
)
body(
    "Explainability in credit risk has gained attention following regulatory requirements. SHAP "
    "(SHapley Additive exPlanations), introduced by Lundberg and Lee (2017), provides theoretically "
    "grounded individual-level feature attribution values. Cross-selling using market basket analysis "
    "(Apriori, FP-Growth) and collaborative filtering has demonstrated strong lift values for financial "
    "product pairs such as Checking-Savings and Credit Card-Personal Loan."
)

# ══════════════════════════════════════════════════════════════════════════════
# 11. CASE STUDY
# ══════════════════════════════════════════════════════════════════════════════
heading("11. CASE STUDY", size=13, color=DARK_BLUE)
heading("11.1 Business Scenario", size=12, color=DARK_BLUE)
body(
    "Consider a retail bank with 2,000 active loan customers. Loan officers currently use "
    "spreadsheet-based credit scoring that takes up to two days per application and provides no "
    "explanation for the decision. The bank's cross-sell conversion rate is below 8% because "
    "recommendations are sent to all customers regardless of their portfolio. The Aether AI system "
    "answers three key business questions in real time:"
)
for q in ["Which loan applicants are most likely to default, and why?",
          "Which customers should be offered which financial products?",
          "How would a macro-economic downturn affect the bank's default exposure?"]:
    bullet(q)

heading("11.2 Example Customer Analysis", size=12, color=DARK_BLUE)
body("Customer ID: C0042", bold=True)
add_table(
    ["Attribute", "Value"],
    [
        ["Age", "38 years"], ["Annual Income", "$45,200"],
        ["Employment Length", "4 years"], ["Credit Score", "572"],
        ["Loan Amount Requested", "$34,000"], ["Owned Products", "Checking, Savings"],
        ["Economic Environment", "Neutral"],
    ]
)
sp()
body("System Output:", bold=True)
add_table(
    ["Metric", "Value", "Interpretation"],
    [
        ["Default Probability", "67.4%", "HIGH RISK — prioritize review"],
        ["Risk Level", "High", "Decline or require collateral"],
        ["Fraud Probability", "3.2%", "Normal — no anomaly"],
        ["Customer Lifetime Value", "$2,112", "Low-medium value segment"],
        ["Top Risk Factor", "Credit Score (572)", "Below safe threshold of 650"],
        ["Second Risk Factor", "Debt-to-Income Ratio", "Loan = 75% of annual income"],
        ["Recommended Products", "Credit Card, Personal Loan", "High portfolio affinity"],
    ]
)
sp()
body("Figure 1: Aether AI — Customer Risk Profile (AI Assessment View)", bold=True, center=True)
add_image("04_customer_risk_profile.png", "Fig 1: Full customer risk dashboard showing default probability gauge, CLV, and fraud score")
sp()
body("Figure 2: SHAP Explainability and Cross-Sell Recommendations", bold=True, center=True)
add_image("05_shap_crosssell_charts.png", "Fig 2: SHAP feature impact chart and personalized cross-sell product recommendations")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 12. DATASET
# ══════════════════════════════════════════════════════════════════════════════
heading("12. DATASET DESCRIPTION", size=13, color=DARK_BLUE)
add_table(
    ["Property", "Value"],
    [
        ["Dataset Name", "Credit Risk Dataset (Banking Customer Data)"],
        ["Total Records", "2,000 customers"],
        ["Total Features", "14 columns (including target)"],
        ["Target Variable", "default (0 = No Default, 1 = Default)"],
        ["Overall Default Rate", "2.3%  (46 defaults out of 2,000)"],
        ["Age Range", "21 to 74 years"],
        ["Income Range", "$20,000 to $500,000  (mean: $86,889)"],
        ["Credit Score Range", "395 to 850  (mean: 646.8)"],
        ["Loan Amount Range", "$1,000 to $100,000  (mean: $21,017)"],
        ["Employment Length", "0 to 39 years  (mean: 19.3 years)"],
    ]
)
sp()
body("Product Ownership Distribution:", bold=True)
add_table(
    ["Product", "Customers Holding", "Percentage"],
    [
        ["Checking Account",   "2,000", "100.0%"],
        ["Savings Account",    "1,431", "71.6%"],
        ["Credit Card",        "963",   "48.2%"],
        ["Auto Loan",          "622",   "31.1%"],
        ["Personal Loan",      "444",   "22.2%"],
        ["Mortgage",           "85",    "4.3%"],
        ["Investment Account", "65",    "3.3%"],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# 13. FEATURE DESCRIPTION
# ══════════════════════════════════════════════════════════════════════════════
heading("13. FEATURE DESCRIPTION", size=13, color=DARK_BLUE)
add_table(
    ["Feature", "Type", "Description", "Role"],
    [
        ["customer_id",       "Integer", "Unique customer identifier",                 "ID"],
        ["age",               "Integer", "Customer age (21–74)",                       "Risk Feature"],
        ["income",            "Float",   "Annual income in USD ($20k–$500k)",          "Risk Feature"],
        ["employment_length", "Integer", "Years of employment (0–39)",                 "Risk Feature"],
        ["credit_score",      "Integer", "Internal credit score (395–850)",            "Risk Feature"],
        ["loan_amount",       "Float",   "Requested loan amount ($1k–$100k)",          "Risk Feature"],
        ["default",           "Binary",  "1 = defaulted, 0 = performing",             "Target"],
        ["Checking",          "Binary",  "1 = holds Checking account",                 "Cross-Sell"],
        ["Savings",           "Binary",  "1 = holds Savings account",                  "Cross-Sell"],
        ["Credit Card",       "Binary",  "1 = holds Credit Card",                      "Cross-Sell"],
        ["Auto Loan",         "Binary",  "1 = holds Auto Loan",                        "Cross-Sell"],
        ["Mortgage",          "Binary",  "1 = holds Mortgage",                         "Cross-Sell"],
        ["Personal Loan",     "Binary",  "1 = holds Personal Loan",                    "Cross-Sell"],
        ["Investment",        "Binary",  "1 = holds Investment account",               "Cross-Sell"],
    ]
)
sp()
heading("13.1 Engineered Features", size=12, color=DARK_BLUE)
add_table(
    ["Derived Feature", "Formula", "Purpose"],
    [
        ["Debt-to-Income Ratio",    "loan_amount / max(income, 1)",                         "Repayment burden"],
        ["Customer Lifetime Value", "(income × 0.05) × (credit_score / 600)",              "Long-term revenue"],
        ["Fraud Anomaly Score",     "High income + very low credit score → spike",         "Flag suspicious profiles"],
        ["SHAP — Credit Score",     "((650 − score) / 100) × 15 × scale",                  "Credit risk driver impact"],
        ["SHAP — Income",           "((60000 − income) / 40000) × 10 × scale",             "Income risk driver impact"],
        ["SHAP — DTI",              "(dti − 0.3) × 30 × scale",                            "Loan burden driver impact"],
        ["SHAP — Employment",       "((5 − emp) / 5) × 5 × scale",                         "Stability driver impact"],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# 14. DATA PREPROCESSING
# ══════════════════════════════════════════════════════════════════════════════
heading("14. DATA PREPROCESSING", size=13, color=DARK_BLUE)
body("14.1  Loading the Dataset", bold=True)
code_block(
    "import pandas as pd\n"
    "data = pd.read_csv('data/banking_data.csv')\n"
    "print(data.shape)          # (2000, 14)\n"
    "print(data['default'].value_counts())\n"
    "# 0    1954\n"
    "# 1      46"
)
body("14.2  Feature Selection", bold=True)
code_block(
    "feature_cols = ['age', 'income', 'employment_length',\n"
    "                'credit_score', 'loan_amount']\n"
    "product_cols = ['Checking', 'Savings', 'Credit Card',\n"
    "                'Auto Loan', 'Mortgage', 'Personal Loan', 'Investment']\n"
    "X = data[feature_cols]\n"
    "y = data['default']"
)
body("14.3  Train-Test Split (80/20, stratified)", bold=True)
code_block(
    "from sklearn.model_selection import train_test_split\n"
    "X_train, X_test, y_train, y_test = train_test_split(\n"
    "    X, y, test_size=0.20, random_state=42, stratify=y\n"
    ")  # Train: 1,600  |  Test: 400"
)
body("14.4  Cross-Sell Affinity Matrix", bold=True)
code_block(
    "cross_sell_rules = {}\n"
    "for product in product_cols:\n"
    "    cross_sell_rules[product] = {}\n"
    "    users_with = data[data[product] == 1]\n"
    "    for other in product_cols:\n"
    "        if product == other: continue\n"
    "        prob = len(users_with[users_with[other] == 1]) / len(users_with)\n"
    "        cross_sell_rules[product][other] = prob"
)

# ══════════════════════════════════════════════════════════════════════════════
# 15. EDA
# ══════════════════════════════════════════════════════════════════════════════
heading("15. EXPLORATORY DATA ANALYSIS", size=13, color=DARK_BLUE)
add_table(
    ["Visualization", "Key Finding"],
    [
        ["Default Distribution",        "2.3% default rate — significant class imbalance"],
        ["Credit Score vs Default",     "Defaulters cluster below credit score 600"],
        ["Income vs Default",           "Defaulters have lower median income ($38k vs $87k)"],
        ["Loan Amount vs Default",      "Defaulters carry disproportionately large loans"],
        ["Debt-to-Income Distribution", "Defaulters show DTI > 0.7 in most cases"],
        ["Employment Length",           "Short employment (0–3 yrs) correlates with default"],
        ["Product Co-ownership",        "Checking-Savings co-ownership rate: 71.6%"],
        ["Risk vs Income Scatter",      "Displayed as interactive scatter plot in dashboard"],
        ["Correlation Heatmap",         "Credit score and default show highest negative correlation"],
        ["Risk Correlation Matrix",     "Bubble chart in Portfolio Analytics tab"],
    ]
)
sp()
body("Figure 3: Portfolio Analytics — Risk Correlation Matrix", bold=True, center=True)
add_image("08_portfolio_analytics.png", "Fig 3: Bubble chart showing correlation strength between income, credit score, loan amount, and default risk")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 16. ML METHODOLOGY
# ══════════════════════════════════════════════════════════════════════════════
heading("16. MACHINE LEARNING METHODOLOGY", size=13, color=DARK_BLUE)
body("The project treats credit default prediction as a binary classification: default=0 means repaid, default=1 means defaulted.")

heading("16.1 Random Forest Classifier", size=12, color=DARK_BLUE)
code_block(
    "from sklearn.ensemble import RandomForestClassifier\n"
    "rf_model = RandomForestClassifier(n_estimators=100, random_state=42)\n"
    "rf_model.fit(X_train, y_train)\n"
    "prob_default = rf_model.predict_proba(X_test)[:, 1]"
)
body(
    "Random Forest builds 100 decision trees on random data subsets. It naturally produces "
    "probability estimates, handles class imbalance via bootstrap sampling, and is interpretable "
    "through feature importances."
)

heading("16.2 SHAP Explainability Approximation", size=12, color=DARK_BLUE)
code_block(
    "def calculate_shap_approximation(cust, prob_default):\n"
    "    base_value = 25.0  # portfolio average default rate\n"
    "    cs_impact  = ((650 - cust['credit_score'])    / 100)   * 15\n"
    "    inc_impact = ((60000 - cust['income'])         / 40000) * 10\n"
    "    emp_impact = ((5 - cust['employment_length']) / 5)     * 5\n"
    "    dti        = cust['loan_amount'] / max(cust['income'], 1)\n"
    "    dti_impact = (dti - 0.3) * 30\n"
    "    total = cs_impact + inc_impact + emp_impact + dti_impact\n"
    "    scale = ((prob_default * 100) - base_value) / total if total else 0\n"
    "    return [\n"
    "        {'feature': 'Credit Score',   'impact': round(cs_impact  * scale, 1)},\n"
    "        {'feature': 'Income',         'impact': round(inc_impact * scale, 1)},\n"
    "        {'feature': 'Debt-to-Income', 'impact': round(dti_impact * scale, 1)},\n"
    "        {'feature': 'Employment',     'impact': round(emp_impact * scale, 1)},\n"
    "    ], base_value"
)

heading("16.3 Macro-Economic Scenario Simulation", size=12, color=DARK_BLUE)
code_block(
    "if   environment == 'Recession':     prob_default = min(prob_default * 1.5, 0.99)\n"
    "elif environment == 'High Interest': prob_default = min(prob_default * 1.2, 0.99)\n"
    "elif environment == 'Booming':       prob_default = max(prob_default * 0.7, 0.01)"
)

heading("16.4 Customer Lifetime Value & Fraud Score", size=12, color=DARK_BLUE)
code_block(
    "# CLV: 5% of income, adjusted by credit quality\n"
    "clv = (cust['income'] * 0.05) * (cust['credit_score'] / 600)\n"
    "\n"
    "# Fraud anomaly rule-based detector\n"
    "fraud_prob = 1.0\n"
    "if cust['income'] > 150000 and cust['credit_score'] < 500: fraud_prob += 45.0\n"
    "if cust['employment_length'] == 0 and cust['loan_amount'] > 50000: fraud_prob += 30.0\n"
    "fraud_prob = min(fraud_prob + random.uniform(0.5, 5.0), 99.9)"
)

# ══════════════════════════════════════════════════════════════════════════════
# 17. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
heading("17. SYSTEM ARCHITECTURE", size=13, color=DARK_BLUE)
for step in [
    "banking_data.csv  (2,000 customers, 14 features)",
    "↓",
    "FastAPI Backend (main.py)  — startup model training",
    "↓",
    "Random Forest Training  +  Cross-Sell Matrix Computation",
    "↓",
    "REST API Endpoints:",
    "    GET  /api/portfolio             — portfolio KPIs",
    "    GET  /api/customers             — customer list",
    "    GET  /api/customers/{id}        — full customer analysis",
    "    POST /api/simulate              — real-time what-if simulation",
    "    GET  /api/portfolio/scatter     — scatter plot data",
    "    GET  /api/customers/{id}/history— transaction timeline",
    "    GET/POST /api/customers/{id}/notes — CRM notes",
    "    POST /api/auth/signup           — user registration",
    "    POST /api/auth/login            — user login",
    "    GET  /api/auth/verify           — session validation",
    "↓",
    "Aether AI Web Dashboard  (HTML + CSS + JavaScript + Chart.js)",
]:
    body(step, indent=True)
sp()
body("Project Structure:", bold=True)
code_block(
    "banking_engine/\n"
    "├── backend/\n"
    "│   ├── main.py              # FastAPI app + ML model\n"
    "│   ├── requirements.txt     # Python dependencies\n"
    "│   ├── users.json           # Hashed user credentials\n"
    "│   ├── notes.json           # CRM notes storage\n"
    "│   └── data/\n"
    "│       └── banking_data.csv # 2,000 customer records\n"
    "├── frontend/\n"
    "│   ├── index.html           # Dashboard HTML\n"
    "│   ├── styles.css           # Glassmorphism UI\n"
    "│   └── app.js               # Dashboard JavaScript\n"
    "├── render.yaml              # Render.com deployment\n"
    "└── start.bat                # One-click local startup"
)

# ══════════════════════════════════════════════════════════════════════════════
# 18. MODEL TRAINING & EVALUATION
# ══════════════════════════════════════════════════════════════════════════════
heading("18. MODEL TRAINING AND EVALUATION", size=13, color=DARK_BLUE)
code_block(
    "from sklearn.metrics import (accuracy_score, precision_score,\n"
    "    recall_score, f1_score, roc_auc_score)\n"
    "\n"
    "y_pred = rf_model.predict(X_test)\n"
    "y_prob = rf_model.predict_proba(X_test)[:, 1]\n"
    "\n"
    "print('Accuracy: ', accuracy_score(y_test, y_pred))\n"
    "print('Precision:', precision_score(y_test, y_pred))\n"
    "print('Recall:   ', recall_score(y_test, y_pred))\n"
    "print('F1-Score: ', f1_score(y_test, y_pred))\n"
    "print('ROC-AUC:  ', roc_auc_score(y_test, y_prob))"
)
add_table(
    ["Metric", "Meaning", "Value"],
    [
        ["Accuracy",  "Overall correct predictions",                     "Report from experiment"],
        ["Precision", "Of predicted defaulters, how many truly defaulted","Report from experiment"],
        ["Recall",    "Of actual defaulters, how many were caught",       "Report from experiment"],
        ["F1-Score",  "Harmonic mean of Precision and Recall",           "Report from experiment"],
        ["ROC-AUC",   "Ability to rank defaulters above non-defaulters",  "Report from experiment"],
    ]
)
body("Replace 'Report from experiment' with your actual results after model training.", italic=True)

# ══════════════════════════════════════════════════════════════════════════════
# 19. RISK CLASSIFICATION
# ══════════════════════════════════════════════════════════════════════════════
heading("19. RISK CLASSIFICATION & EARLY WARNING MECHANISM", size=13, color=DARK_BLUE)
add_table(
    ["Default Probability", "Risk Level", "Color", "Suggested Action"],
    [
        ["0% – 20%",   "Low",    "Green", "Standard processing — routine review"],
        ["21% – 50%",  "Medium", "Amber", "Request additional income verification"],
        ["51% – 100%", "High",   "Red",   "Decline, require co-signer, or reduce amount"],
    ]
)
sp()
body("Figure 4: What-If Simulator and Decision Engine", bold=True, center=True)
add_image("06_whatif_simulator.png", "Fig 4: What-If Simulator with loan/income sliders and macro-economic environment selector")
sp()
body("Figure 5: Decision Engine Output", bold=True, center=True)
add_image("07_decision_engine.png", "Fig 5: Decision engine showing APPROVED / MANUAL REVIEW / DECLINED outcome")

# ══════════════════════════════════════════════════════════════════════════════
# 20. CROSS-SELL ENGINE
# ══════════════════════════════════════════════════════════════════════════════
heading("20. CROSS-SELL RECOMMENDATION ENGINE", size=13, color=DARK_BLUE)
body(
    "The cross-sell engine uses market basket analysis (conditional probability scoring) to "
    "recommend the top-3 products a customer is most likely to need based on their existing portfolio."
)
code_block(
    "for un in unowned_products:\n"
    "    score = sum(cross_sell_rules[owned].get(un, 0) for owned in owned_products)\n"
    "    if un == 'Mortgage'   and income > 80000:  score += 1.0\n"
    "    if un == 'Investment' and income > 100000: score += 0.8\n"
    "    confidence = min(score / max(len(owned_products), 1) * 100, 95)\n"
    "\n"
    "top_recs = sorted(recs, key=lambda x: x['confidence'], reverse=True)[:3]"
)
body("Example output for a customer who owns Checking + Savings:", bold=True)
add_table(
    ["Rank", "Recommended Product", "Confidence", "Reason"],
    [
        ["1", "Credit Card",   "62.3%", "High co-ownership rate among Savings holders"],
        ["2", "Auto Loan",     "44.7%", "Common next product for Savings holders"],
        ["3", "Personal Loan", "31.2%", "Frequent combination after Credit Card"],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# 21. IMPLEMENTATION
# ══════════════════════════════════════════════════════════════════════════════
heading("21. SYSTEM IMPLEMENTATION", size=13, color=DARK_BLUE)
add_table(
    ["Component", "Technology"],
    [
        ["Language",       "Python 3.12"],
        ["Web Framework",  "FastAPI 0.115+"],
        ["ML Library",     "scikit-learn (RandomForestClassifier)"],
        ["Data",           "Pandas, NumPy"],
        ["Web Server",     "Uvicorn (ASGI)"],
        ["Frontend",       "HTML5, CSS3, Vanilla JavaScript"],
        ["Charts",         "Chart.js (CDN)"],
        ["Authentication", "hashlib SHA-256 + secrets module"],
        ["Deployment",     "Render.com (render.yaml)"],
        ["Startup",        "start.bat (Windows batch file)"],
    ]
)
sp()
body("requirements.txt:", bold=True)
code_block("fastapi\nuvicorn[standard]\npandas\nnumpy\nscikit-learn\npydantic\naiofiles")

# ══════════════════════════════════════════════════════════════════════════════
# 22. DASHBOARD
# ══════════════════════════════════════════════════════════════════════════════
heading("22. WEB DASHBOARD — AETHER AI", size=13, color=DARK_BLUE)
add_table(
    ["Screen", "Key Components"],
    [
        ["Login / Signup",      "Glassmorphism auth overlay, sample credentials badge, tab switcher"],
        ["AI Assessment",       "Risk gauge, CLV, Fraud score, SHAP chart, Scatter plot, Cross-sell, History, CRM notes, Simulator, Decision Engine"],
        ["Portfolio Analytics", "Bubble correlation matrix"],
        ["New Analysis",        "6-field custom input form with sample data, full results panel"],
    ]
)
sp()
body("Figure 6: Login Screen — Aether AI", bold=True, center=True)
add_image("01_login_screen.png", "Fig 6: Glassmorphism login screen with sample credentials badge")
sp()
body("Figure 7: Login with Demo Credentials", bold=True, center=True)
add_image("02_login_credentials.png", "Fig 7: Login form filled with demo credentials before submission")
sp()
body("Figure 8: Main Dashboard after Login", bold=True, center=True)
add_image("03_dashboard_overview.png", "Fig 8: Main dashboard showing empty AI Assessment view after successful login")
sp()
body("Figure 9: New Analysis — Custom Customer Input Form", bold=True, center=True)
add_image("09_new_analysis_form.png", "Fig 9: New Analysis tab with pre-filled sample data form")
sp()
body("Figure 10: New Analysis — Risk Results", bold=True, center=True)
add_image("10_custom_results_risk.png", "Fig 10: Custom analysis results showing risk gauge, CLV, fraud score")
sp()
body("Figure 11: New Analysis — SHAP and Risk Summary", bold=True, center=True)
add_image("11_custom_shap_summary.png", "Fig 11: SHAP explainability chart and risk summary table for custom customer")
sp()
body("Figure 12: New Analysis — Decision Engine Output", bold=True, center=True)
add_image("12_custom_decision_output.png", "Fig 12: Decision engine output after running analysis on custom customer")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 23. EXPECTED RESULTS
# ══════════════════════════════════════════════════════════════════════════════
heading("23. EXPECTED RESULTS", size=13, color=DARK_BLUE)
for item in [
    "A trained Random Forest model computing individual default probabilities.",
    "SHAP-approximated risk driver explanations for every customer.",
    "Top-3 personalized product recommendations per customer.",
    "Customer Lifetime Value estimates for revenue prioritization.",
    "Fraud anomaly scores to flag suspicious applications.",
    "Macro-economic scenario simulation results.",
    "A fully functional Aether AI dashboard with secure multi-user authentication.",
    "A reusable API with 10+ endpoints serving all dashboard functions.",
]:
    bullet(item)
body("Replace all illustrative values with actual model results after training.", italic=True)

# ══════════════════════════════════════════════════════════════════════════════
# 24. ADVANTAGES
# ══════════════════════════════════════════════════════════════════════════════
heading("24. ADVANTAGES", size=13, color=DARK_BLUE)
for item in [
    "Unified platform for credit risk assessment and cross-sell optimization.",
    "Individual probability estimates more informative than binary classifications.",
    "SHAP explainability supports regulatory compliance.",
    "Real-time simulation stress-tests the portfolio under recession scenarios.",
    "Market basket analysis improves cross-sell conversion by targeting relevant products.",
    "CLV scores prioritize high-value customers for premium product offers.",
    "Fraud detection adds an additional layer of application security.",
    "FastAPI backend is highly performant and supports concurrent API calls.",
]:
    bullet(item)

# ══════════════════════════════════════════════════════════════════════════════
# 25. LIMITATIONS
# ══════════════════════════════════════════════════════════════════════════════
heading("25. LIMITATIONS", size=13, color=DARK_BLUE)
for item in [
    "Dataset is synthetically generated; real-world data may have more complex distributions.",
    "2.3% default rate causes class imbalance; minority class recall may be limited.",
    "SHAP values are approximated, not computed from Shapley game theory.",
    "In-memory session storage (SESSIONS dict) resets when the server restarts.",
    "CRM notes and user data in JSON files are not suitable for large-scale production.",
    "No real-time data ingestion; model is trained once at startup on static CSV data.",
]:
    bullet(item)

# ══════════════════════════════════════════════════════════════════════════════
# 26. FUTURE SCOPE
# ══════════════════════════════════════════════════════════════════════════════
heading("26. FUTURE SCOPE", size=13, color=DARK_BLUE)
for item in [
    "Replace SHAP approximation with exact SHAP values using the shap Python library.",
    "Add XGBoost and LightGBM as alternative models and compare metrics.",
    "Implement SMOTE or class-weight balancing for imbalanced data.",
    "Replace JSON storage with PostgreSQL for production-grade persistence.",
    "Replace in-memory sessions with Redis-backed JWT authentication.",
    "Add customer segmentation using K-Means clustering.",
    "Integrate a recommendation model (collaborative filtering) for cross-sell.",
    "Add real-time data ingestion from a banking transaction stream.",
    "Implement model drift monitoring and periodic retraining pipelines.",
    "Deploy automated alerts via email when high-risk customers are detected.",
]:
    bullet(item)

# ══════════════════════════════════════════════════════════════════════════════
# 27. CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
heading("27. CONCLUSION", size=13, color=DARK_BLUE)
body(
    "The Banking Credit Default Risk and Cross-Sell Engine is a practical AIML case study "
    "demonstrating how financial customer data can be transformed into an explainable, real-time, "
    "and actionable intelligence platform. The system addresses two commercially significant "
    "challenges in retail banking: preventing loan default losses and maximizing revenue through "
    "targeted product cross-selling."
)
body(
    "The project connects data preprocessing, Random Forest classification, SHAP-approximated "
    "explainability, market basket analysis, Customer Lifetime Value estimation, fraud anomaly "
    "detection, macro-economic scenario simulation, and a production-quality glassmorphism web "
    "dashboard — all served through a FastAPI REST API. Secure login and signup ensures only "
    "authorized bank staff can access the system."
)
body(
    "The strongest final implementation should replace all illustrative metric values with actual "
    "experimental results, clearly document the dataset source and preprocessing steps, and include "
    "screenshots of the working Aether AI dashboard."
)

# ══════════════════════════════════════════════════════════════════════════════
# 28. REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
heading("28. REFERENCES", size=13, color=DARK_BLUE)
for ref in [
    "Scikit-learn documentation — RandomForestClassifier, model_selection, metrics. https://scikit-learn.org",
    "Lundberg, S. M., & Lee, S. I. (2017). A Unified Approach to Interpreting Model Predictions. NeurIPS 2017.",
    "FastAPI documentation — FastAPI framework for building REST APIs with Python. https://fastapi.tiangolo.com",
    "IBM Resources — Credit risk management and machine learning in banking. https://www.ibm.com/topics/credit-risk",
    "Chart.js documentation — Interactive charts for web dashboards. https://www.chartjs.org",
    "Han, J., Kamber, M., & Pei, J. — Data Mining: Concepts and Techniques (Market Basket Analysis).",
    "Baesens, B. et al. — Credit Risk Analytics: Measurement Techniques, Applications and Examples.",
    "Thomas, L. C. — Consumer Credit Models: Pricing, Profit and Portfolios.",
    "Render.com — Cloud deployment of Python web applications. https://render.com/docs",
    "python-docx documentation — Creating Word documents with Python. https://python-docx.readthedocs.io",
]:
    bullet(ref)

# ── Save ──────────────────────────────────────────────────────────────────────
out = "C:/Users/sait9/OneDrive/Desktop/banking_engine/IBM_CaseStudy_T_Sai_Santosh_Sri_Ram.docx"
doc.save(out)
import os
kb = os.path.getsize(out) // 1024
print(f"Saved: {out}")
print(f"Size : {kb} KB")
